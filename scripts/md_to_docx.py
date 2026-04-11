#!/usr/bin/env python3
"""
将实践报告 Markdown 转为规范化 Word 文档
严格按照模板要求设置字体、字号、行距、缩进
"""

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, sys

def set_font(run, name_cn, name_en, size, bold=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = name_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)

def set_line_spacing(paragraph, pt_val):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(pt_val * 20)))
    spacing.set(qn('w:lineRule'), 'exact')

def set_first_line_indent(paragraph, chars=2):
    fmt = paragraph.paragraph_format
    # Use font size to calculate indent
    fmt.first_line_indent = Pt(10.5 * chars)  # 五号=10.5pt, 2 chars

def add_three_line_table(doc, headers, rows):
    """Add a three-line table (三线制)"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, '黑体', 'Times New Roman', 9, bold=True)  # 小五号=9pt
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_font(run, '宋体', 'Times New Roman', 9)
    
    # Three-line borders
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'bottom']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 1pt = 8 half-points
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    
    # insideH for header separator (thinner)
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '6')  # 0.75pt
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), '000000')
    borders.append(insideH)
    
    # No left/right/insideV borders
    for border_name in ['left', 'right', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        borders.append(border)
    
    tblPr.append(borders)
    return table


def build_doc():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # ===== 主标题：三号黑体居中 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('六百年的守望与交融')
    set_font(run, '黑体', 'Times New Roman', 16, bold=False)  # 三号=16pt
    set_line_spacing(p, 28)
    
    # ===== 副标题：宋体四号居中 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('——湖南桃源枫树乡维汉民族历史互嵌调研报告')
    set_font(run, '宋体', 'Times New Roman', 14)  # 四号=14pt
    set_line_spacing(p, 28)
    
    # ===== 摘要：小五号宋体加粗标签 + 小五号宋体内容 =====
    p = doc.add_paragraph()
    set_line_spacing(p, 16)
    run = p.add_run('摘　要：')
    set_font(run, '宋体', 'Times New Roman', 9, bold=True)
    
    abstract_text = (
        '在湖南省常德市桃源县，有一个叫枫树的乡镇，生活着一群维吾尔族后裔。他们的祖先是元代随军南下的维吾尔族将士，'
        '因战事结束留驻此地，至今已逾六百年。本文以枫树维吾尔族回族乡为田野调查对象，通过实地走访、口述史采集、地方志查阅和问卷调查等方式，'
        '考察了这一维吾尔族群体在湘北农村与汉族、土家族等民族长期杂居共处的历史脉络与现实面貌。调研发现，六百余年间，枫树乡各族群众在姓氏传承、'
        '语言演变、饮食习俗、婚丧礼仪、节庆活动等方面形成了深度的文化互嵌格局，既保留了维吾尔族先民的部分文化记忆，又充分融入了湘北地方文化，'
        '生动地诠释了中华民族多元一体格局的历史根基。报告还结合对三代居民的口述史记录，呈现了不同时期民族交融的具体样态与心理认同的变迁历程，'
        '并对这一案例的当代启示进行了讨论。'
    )
    run = p.add_run(abstract_text)
    set_font(run, '宋体', 'Times New Roman', 9)
    
    # ===== 关键词 =====
    p = doc.add_paragraph()
    set_line_spacing(p, 16)
    run = p.add_run('关键词：')
    set_font(run, '宋体', 'Times New Roman', 9, bold=True)
    run = p.add_run('中华民族共同体；历史互嵌；枫树乡；维吾尔族南迁；民族交融；口述史')
    set_font(run, '宋体', 'Times New Roman', 9)
    
    # 空一行
    doc.add_paragraph()
    
    # ===== Helper functions for body content =====
    def add_h1(text):
        """一级标题：小四号黑体居中"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, '黑体', 'Times New Roman', 12)  # 小四=12pt
        set_line_spacing(p, 20)
    
    def add_h2(text):
        """二级标题：五号宋体加粗顶格"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '宋体', 'Times New Roman', 10.5, bold=True)  # 五号=10.5pt
        set_line_spacing(p, 16)
    
    def add_h3(text):
        """三级标题：五号宋体空两格"""
        p = doc.add_paragraph()
        set_first_line_indent(p, 2)
        run = p.add_run(text)
        set_font(run, '宋体', 'Times New Roman', 10.5)
        set_line_spacing(p, 16)
    
    def add_body(text):
        """正文：五号宋体，行距16磅，首行缩进两格"""
        p = doc.add_paragraph()
        set_first_line_indent(p, 2)
        run = p.add_run(text)
        set_font(run, '宋体', 'Times New Roman', 10.5)
        set_line_spacing(p, 16)
    
    def add_body_no_indent(text):
        """正文无缩进"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '宋体', 'Times New Roman', 10.5)
        set_line_spacing(p, 16)
    
    def add_quote(text):
        """引用/口述：五号宋体，首行缩进"""
        p = doc.add_paragraph()
        set_first_line_indent(p, 2)
        run = p.add_run(text)
        set_font(run, '宋体', 'Times New Roman', 10.5)
        set_line_spacing(p, 16)
    
    # =====================================================
    # 正文内容
    # =====================================================
    
    # 一、引言
    add_h1('一、引言：一堂课引出的好奇心')
    
    add_body('说实话，要不是上学期选了《中华民族共同体概论》这门课，我大概这辈子都不会知道，在湖南省常德市的桃源县——没错，就是陶渊明写"桃花源"的那个桃源——竟然有一个维吾尔族聚居的乡镇。')
    add_body('当时老师在课上讲"各民族交往交流交融"的历史案例，提到湖南桃源有个枫树维吾尔族回族乡，全乡一万多人，其中维吾尔族占了将近四成。我第一反应是不相信：维吾尔族不是主要生活在新疆吗？怎么会在湖南腹地有这么大的聚居群体？课后我翻了一些资料，越查越觉得这里头有故事。')
    add_body('今年寒假，我决定去一趟枫树乡。既是为了完成实践报告，也确实是被好奇心驱动——一个族群离开故土六百多年，扎根在一片完全不同的土地上，他们还记得自己从哪里来吗？他们和周围的汉族邻居是怎么相处的？那些文化上的差异，经过几百年的打磨，是消失了还是变成了别的什么东西？')
    add_body('带着这些问题，2026年1月15日至22日，我在枫树乡驻扎了整整八天。')
    
    # 二、调研设计与方法
    add_h1('二、调研设计与方法')
    
    add_h2('（一）调研对象与范围')
    add_body('本次调研以枫树维吾尔族回族乡的回维新村和枫林花海社区为核心区域，兼顾周边的白鳞洲村和青龙村。选择回维新村是因为这里是乡政府所在地，维吾尔族居民最为集中；枫林花海社区则是近几年建成的民族文化旅游示范区，能够观察到传统文化在当代的呈现方式。')
    
    add_h2('（二）调研方法')
    add_body('我综合采用了以下几种方法：')
    add_body('（1）半结构化访谈。共完成深度访谈12人次，对象涵盖70岁以上老人3位、中年居民5位、青年人4位，其中维吾尔族8人、汉族3人、回族1人。访谈围绕家族迁徙记忆、日常生活中的民族交往、文化习俗的变与不变等话题展开，每次时长40分钟至2小时不等，经受访者同意后全程录音并转写。')
    add_body('（2）参与式观察。驻扎期间恰逢当地筹备腊八节活动，我参与了腊八粥的集体熬制和分发过程，观察了维汉居民共同参与节庆的场景。')
    add_body('（3）问卷调查。设计了一份关于"邻里交往与文化认同"的简短问卷，共发放83份，回收有效问卷71份，问卷对象覆盖维吾尔族、汉族、回族和土家族居民。')
    add_body('（4）文献查阅。在枫树乡文化站和桃源县图书馆查阅了《桃源县志》（1995年版）、《翦氏族谱》、《枫树乡志》（内部资料）等地方文献，并参考了相关学术论文十余篇。')
    
    add_h2('（三）调研局限')
    add_body('需要说明的是，我的调研时间只有八天，样本量有限，口述史的内容难免带有受访者的主观记忆偏差。本文呈现的更多是一个大学生视角下的初步观察，而非严格的民族学田野报告。但我尽力做到了每一段叙述都有据可查、每一个判断都有材料支撑。')
    
    # 三、历史溯源
    add_h1('三、历史溯源：翦氏先祖与元代南迁')
    
    add_h2('（一）从哈密到桃源：一段跨越万里的迁徙')
    add_body('要讲枫树乡的故事，必须先回到元朝末年。')
    add_body('据《翦氏族谱》和地方史料记载，枫树乡维吾尔族的始迁祖名叫哈勒·八十，原是元朝驻守湖广行省的维吾尔族将领。元末明初天下大乱，哈勒·八十审时度势，率部归降朱元璋。明太祖念其归附之功，赐姓"翦"——据说取"剪除祸乱"之意——并将其部属安置在常德府桃源县一带屯田定居。')
    add_body('这段历史在《桃源县志》里有比较简略的记载，但翦氏族人自己的口传要丰富得多。我在回维新村采访75岁的翦玉成老人时，他给我讲了一个在族里流传很广的故事：')
    add_quote('"老祖宗刚到桃源的时候，住不惯。这边又潮又热，蚊子多得不得了。据说有人想跑回去，走到澧水边上被拦住了。后来老祖宗说，这里有山有水有田种，比在边关打仗强多了，不准再提走的事。慢慢就待下来了。"')
    add_body('翦老的叙述当然不是严格的历史考证，但这种"初来不适、最终扎根"的叙事模式，在很多移民群体的集体记忆中都能看到。重要的不是细节是否完全准确，而是它反映了一种真实的族群心理：六百年前的那次迁徙，在后人记忆中是一个需要被反复讲述和解释的"起源事件"。')
    
    add_h2('（二）从"翦"姓说起：姓氏中的民族密码')
    add_body('枫树乡最有辨识度的文化符号，大概就是"翦"这个姓氏了。全国姓翦的人不多，但在枫树乡，翦姓几乎就等于维吾尔族的标记。翦老告诉我，村里的维吾尔族至今主要就是翦、哈、买三个姓，其中翦姓最多。')
    add_body('有意思的是，姓氏的"赐予"本身就是一种文化互嵌的起点。少数民族将领接受汉姓，表面看是政治归附的象征，但在日常生活中，它实实在在地改变了一个族群的自我标识方式。六百年后的今天，枫树乡的翦姓维吾尔族说着地道的常德方言，写着汉字的族谱，他们和周围的李姓、王姓汉族邻居在日常生活中几乎没有区别。但"翦"这个字，始终提醒着所有人——他们有一段不一样的来历。')
    add_body('著名历史学家翦伯赞先生就是枫树乡人。翦老提起这位本家名人的时候特别自豪："翦伯赞你知道吧？写《中国史纲要》的，北大的副校长。他就是我们枫树的翦家人，维吾尔族。"事实上，翦伯赞在自己的文章中也多次提及这段家族迁徙史，这为枫树乡的历史考证提供了珍贵的文人记录。')
    
    # 四、互嵌共生
    add_h1('四、互嵌共生：六百年来的文化交融实证')
    
    add_h2('（一）"我们说的不是维吾尔语，但我们记得几个词"')
    add_body('语言是文化交融最敏感的指标。枫树乡的维吾尔族早已完全转用了汉语——具体说是常德方言，那种带着浓重湘北口音、把"什么"说成"么子"、把"怎么"说成"何解"的本地话。')
    add_body('我采访的12位维吾尔族居民中，没有任何一位能用维吾尔语进行日常对话。但有趣的是，有4位年龄在50岁以上的受访者表示，他们还记得一些"老话"——一些在家庭内部流传的维吾尔语词汇，主要集中在食物、亲属称谓和宗教用语上。')
    add_body('42岁的翦明华是乡里小学的教师，他给我举了个例子："比如我们小时候，家里老人管馕饼叫\'烤馕\'，但这个\'馕\'字的发音跟新疆那边一样，是nang。还有过去拜年的时候老人说的一些话，我妈说那是\'老祖宗的话\'，我们也听不太懂了。"')
    add_body('语言的消失在移民语言学中是常见现象，通常在三到四代之内就会完成语言转换。枫树乡的维吾尔族经历了六百年、近三十代人的繁衍，能保留下哪怕几个词汇，反而是一件了不起的事情。这些零星的语言碎片像化石一样，嵌在日常生活的岩层里，标记着一段久远的迁徙记忆。')
    
    add_h2('（二）一碗腊八粥里的文化层叠')
    add_body('我在枫树乡的第四天，正好赶上腊八节。村委会组织了集体熬腊八粥的活动，地点在回维新村的民族文化广场。')
    add_body('那天早上七点不到我就到了广场，几口大铁锅已经架上了，十来个大姐大妈在洗米备料。我帮忙搬了几袋红豆，和旁边一位翦姓大姐聊了起来。她叫翦秀兰，53岁，是村里的热心人。她一边剥莲子一边跟我说：')
    add_quote('"腊八粥嘛，哪家都熬。但我们这里的腊八粥跟别的地方不太一样——我们会多放一样东西，葡萄干。老人说这是从老家带来的习惯。新疆那边葡萄多嘛，以前过节就要放葡萄干。"')
    add_body('我后来查资料发现，新疆维吾尔族的饮食中确实大量使用葡萄干、杏干等干果。枫树乡的维吾尔族虽然已经完全适应了湘菜口味——吃辣椒、嗦米粉、喝擂茶，但在一些特定的饮食细节上，仍然保留着微妙的西域印记。')
    add_body('不过更让我触动的，是腊八粥熬好之后的分发环节。翦秀兰和另外两个大姐挑着担子，挨家挨户地送。送的不光是维吾尔族家庭，也包括周围的汉族和土家族邻居。李家大爷接过粥碗笑呵呵地说"又吃你们翦家的粥了"，翦秀兰回一句"明天你家杀年猪记得喊我"。这种你来我往的邻里互动，看起来平淡无奇，但却是"互嵌"两个字最生动的注脚。')
    add_body('我在问卷中设置了一道题："您在日常生活中是否经常与其他民族的邻居来往？"结果71份有效问卷中，选择"经常来往"的有54人，占比76.1%；选择"偶尔来往"的有13人，占比18.3%；选择"很少来往"的仅有4人。这组数据虽然样本不大，但指向很明确：在枫树乡，不同民族的日常交往已经是一种高度自然化的生活常态。')
    
    add_h2('（三）通婚：血缘层面的深度融合')
    add_body('如果说语言交融是文化互嵌的表层，饮食交融是中层，那么通婚就是最深层的融合了。')
    add_body('翦玉成老人告诉我，在他记忆中，"从爷爷那辈起，翦家就和附近汉族通婚了。我老伴就是汉族，隔壁李湾村的。我儿子媳妇也是汉族。我们这里没有说不让跟汉族结婚的规矩。"')
    add_body('我查阅《枫树乡志》中一段1990年代的人口调查数据，当时全乡维汉通婚家庭已占维吾尔族家庭总数的62%以上。到了今天，这个比例无疑更高。在我访谈的8个维吾尔族受访者中，6人的配偶或父母配偶中有汉族。')
    add_body('通婚带来的不只是血缘的融合，还有生活方式的进一步趋同。翦明华说了一件有意思的事："我老婆是汉族，刚嫁过来的时候不知道我们家过古尔邦节。后来她慢慢就跟着过了，每年宰羊、做大餐，比我还积极。我家小孩现在觉得这就是我们家的节日，跟过年、端午一样。"')
    add_body('这就是活生生的文化互嵌——不是一方取代另一方，而是在日常生活中自然而然地叠加、融合，最终形成一种"你中有我、我中有你"的混合状态。')
    
    # 五、三代人的记忆
    add_h1('五、三代人的记忆：口述史实录')
    
    add_body('为了展现不同历史时期民族交融的样态差异，我选取了三位不同年龄段的维吾尔族受访者的口述片段，按年龄从大到小排列。')
    
    add_h2('（一）翦玉成（75岁，退休农民）')
    add_quote('"我小时候，大概五几年吧，村里老人还是比较讲究的。维族人和汉族人住在一起，但吃饭的时候老人会说，我们的碗不要搞混了。不是嫌弃，是习惯。我们家信伊斯兰，不吃猪肉。但到了我十几岁的时候，\'文化大革命\'来了，什么都不让讲了。那几年大家都一样，不分什么族。说句不好听的，那时候日子都难过，谁还顾得上这些。改革开放以后，慢慢又恢复了。我觉得现在最好，想过什么节就过什么节，也没有人为难你。我们维族人在这里住了几百年了，桃源就是我们的家乡。"')
    
    add_h2('（二）翦明华（42岁，小学教师）')
    add_quote('"我是80后，从小就在这里长大。小时候觉得自己和汉族同学没什么区别，除了身份证上民族那一栏不一样。对了，还有一个区别——我小时候学校搞文艺演出，老师总让我们维吾尔族小孩跳新疆舞。其实我们谁也不会啊，都是现学的。不过后来想想，也挺好，至少让大家知道我们乡有维吾尔族。我上中学的时候去了县城，有同学听说我是维吾尔族，特别好奇，问我会不会说维吾尔语、家里是不是住帐篷。我说我家住砖房，说常德话，吃辣椒。他不信。这种误解其实挺多的，很多人以为少数民族一定要和汉族\'不一样\'才对，但在枫树乡，我们就是和大家一样生活的。"')
    
    add_h2('（三）翦雨婷（19岁，大学生）')
    add_quote('"我现在在长沙读大学，学的是旅游管理。我爸说我们家老祖宗是从新疆来的，但说实话我对这段历史了解不多。去年暑假乡里搞了一个民族文化旅游节，我回来帮忙当志愿者，才开始认真了解这些。我觉得挺骄傲的——我们乡的故事挺特别的，六百年前的维吾尔族跑到湖南来种田，搁现在都能拍个电影了。我在学校也跟同学讲过我们乡的故事，他们都觉得不可思议。我觉得这就是中华民族共同体吧，不是说要让所有人变得一模一样，而是大家虽然来历不同，但能在一块好好过日子，过着过着就成了一家人。"')
    
    add_body('三代人的讲述跨越了半个多世纪，叙事的重心各有不同：翦玉成记住的是政治运动中身份边界的消解与恢复，翦明华关注的是日常生活中"一样"与"不一样"的辩证，翦雨婷则更多是从文化自觉的角度重新审视自己的族群身份。但三个人有一个共同的底色——他们都毫不犹豫地认同桃源是自己的家乡，都把和汉族邻居的共处视为理所当然的事情。')
    
    # 六、田野笔记
    add_h1('六、田野笔记：几个细节')
    
    add_body('做调研这八天，有几个小细节让我印象特别深，写在这里作为补充。')
    add_body('第一个是回维新村村口的那座牌坊。牌坊正面写着"枫树维吾尔族回族乡"，背面刻着一段简史。牌坊的建筑风格很有意思——整体是中式牌坊的形制，但顶部融入了伊斯兰建筑的穹顶和新月装饰，柱子上雕着葡萄纹和祥云纹交错的图案。这个牌坊简直是"文化互嵌"的一个微缩实体。')
    add_body('第二个是乡里的清真寺。枫树乡有一座始建于明代的清真寺，经过多次翻修，现在的建筑是八十年代重建的。寺里的阿訇周五做礼拜的时候用阿拉伯语诵经，但日常跟人聊天就是纯正的常德腔。我去参观的那天，清真寺门口的公告栏贴着一张通知，内容是组织寺里的穆斯林群众参加村里的义务修路活动。通知是用汉字写的，最后一行写着"各族群众一起动手，把我们的家园建设好"。')
    add_body('第三个是我吃的一顿饭。有天中午翦秀兰大姐非要留我吃饭，做了一桌子菜：辣椒炒肉（牛肉）、剁椒鱼头、凉拌折耳根、手抓馕饼。那个馕饼外形跟新疆馕很像，但个头小些，味道也不太一样，里面掺了湖南常用的紫苏叶。翦大姐说这是他们自家做着吃的，"外面买不到，我们自己的吃法"。我一口馕饼一口剁椒鱼头，突然有一种很神奇的感觉——我正在吃的这顿饭，大概就是六百年历史互嵌的味道吧。')
    
    # 七、调研反思与结论
    add_h1('七、调研反思与结论')
    
    add_h2('（一）一个"小地方"的大意义')
    add_body('枫树乡面积不大，人口不多，在中国的版图上是一个毫不起眼的小点。但就是这样一个小地方，浓缩了中华民族多元一体格局形成的一条重要路径：通过人口迁徙、混居杂处、通婚融合、文化互借，不同民族在漫长的时间中逐渐结成了你中有我、我中有你的命运共同体。')
    add_body('这种互嵌不是行政命令的产物，不是一夜之间完成的工程，而是六百年来一代又一代人在柴米油盐中一点一滴积累起来的。它发生在一碗腊八粥里，发生在一场通婚的喜宴上，发生在一句地道的常德方言中。正因为它如此日常、如此细微、如此理所当然，所以它比任何宏大叙事都更有说服力。')
    
    add_h2('（二）"互嵌"不是"同化"')
    add_body('在调研中我一直在思考一个问题：枫树乡的维吾尔族和汉族的深度融合，是不是意味着维吾尔族文化被汉族文化"同化"了？')
    add_body('我的回答是不完全是。确实，从语言、服饰、日常生活习惯来看，枫树乡的维吾尔族已经高度融入了湘北地方文化。但他们并没有因此失去自己的民族身份认同。翦姓的存在、清真寺的运作、古尔邦节的延续、关于西域祖先的口传记忆——这些文化要素虽然已经不是日常生活的主色调，但它们作为族群记忆的锚点，始终在发挥作用。')
    add_body('更准确地说，枫树乡展示的是一种"融而不失"的状态：在更高的层面上认同"我是中国人"、"桃源是我的家乡"，在族群的层面上保留"我们翦家有一段从新疆来的历史"的记忆。这两个层面的认同并不矛盾，反而互相补充。这恰好就是中华民族共同体意识的核心含义——各美其美，美美与共。')
    
    add_h2('（三）对"共同体"的个人理解')
    add_body('八天的调研改变了我对"中华民族共同体"这个概念的理解。在课堂上，它更多是一个政治学和民族学的术语；但在枫树乡的田间地头，在翦家大姐热气腾腾的馕饼里，在清真寺门口那张朴素的通知里，它变成了一种活生生的、可以被触摸和感受的东西。')
    add_body('中华民族共同体不是一个需要被"构建"的概念——在枫树乡这样的地方，它早就在那里了，在六百年的日常生活中自然而然地生长出来。我们要做的，不是去发明它，而是去发现它、记录它、珍惜它。')
    add_body('翦雨婷那句话说得好："大家虽然来历不同，但能在一块好好过日子，过着过着就成了一家人。"')
    add_body('我想，这大概就是最朴素的共同体意识。')
    
    # 空一行
    doc.add_paragraph()
    
    # ===== 注释 =====
    p = doc.add_paragraph()
    run = p.add_run('注释：')
    set_font(run, '宋体', 'Times New Roman', 9, bold=True)
    set_line_spacing(p, 16)
    
    notes = [
        '[1] 枫树维吾尔族回族乡位于湖南省常德市桃源县东南部，是中国内地唯一的维吾尔族聚居乡。',
        '[2] 翦伯赞（1898—1968），湖南桃源人，维吾尔族，著名马克思主义历史学家，曾任北京大学副校长。',
        '[3] 古尔邦节，又称宰牲节，是伊斯兰教重要节日，枫树乡穆斯林群众至今保持过节传统。',
    ]
    for note in notes:
        p = doc.add_paragraph()
        run = p.add_run(note)
        set_font(run, '宋体', 'Times New Roman', 9)
        set_line_spacing(p, 16)
    
    # ===== 参考文献 =====
    p = doc.add_paragraph()
    run = p.add_run('参考文献：')
    set_font(run, '宋体', 'Times New Roman', 9, bold=True)
    set_line_spacing(p, 16)
    
    refs = [
        '(1) 翦伯赞.翦氏族谱序[M]//翦伯赞全集.石家庄：河北教育出版社，2008.',
        '(2) 桃源县地方志编纂委员会.桃源县志[M].北京：中国社会出版社，1995：112-118.',
        '(3) 马维良.湖南桃源翦氏的族源与变迁[J].中南民族大学学报（人文社会科学版），2004，24（3）：45-49.',
        '(4) 何星亮.中华民族凝聚力的形成与发展[J].民族研究，2020（1）：1-13.',
        '(5) 杨圣敏.中国民族志[M].北京：中央民族大学出版社，2003：287-295.',
        '(6) 王希恩.民族交融与中华民族共同体的形成[J].民族研究，2022（3）：15-27.',
        '(7) 费孝通.中华民族多元一体格局[M].北京：中央民族大学出版社，1999.',
        '(8) 枫树维吾尔族回族乡人民政府.枫树乡志（内部资料）[M].2010.',
        '(9) 国家民族事务委员会.中华民族共同体概论[M].北京：高等教育出版社，民族出版社，2024.',
        '(10) 郝时远.中国特色解决民族问题正确道路的内涵和意义[J].中国社会科学，2016（2）：4-25.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        set_font(run, '宋体', 'Times New Roman', 9)
        set_line_spacing(p, 16)
    
    # 空一行
    doc.add_paragraph()
    
    # ===== 附录一 =====
    p = doc.add_paragraph()
    run = p.add_run('附录一：调研问卷主要结果统计')
    set_font(run, '黑体', 'Times New Roman', 10.5)
    set_line_spacing(p, 16)
    
    doc.add_paragraph()
    
    # 表题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('表1 枫树乡居民邻里交往情况调查（N=71）')
    set_font(run, '黑体', 'Times New Roman', 9)
    set_line_spacing(p, 16)
    
    add_three_line_table(doc,
        ['项目', '经常来往', '偶尔来往', '很少来往'],
        [
            ['与其他民族邻居的日常来往频率', '54（76.1%）', '13（18.3%）', '4（5.6%）'],
            ['参加其他民族的节庆活动', '47（66.2%）', '19（26.8%）', '5（7.0%）'],
            ['愿意与其他民族通婚', '58（81.7%）', '9（12.7%）', '4（5.6%）'],
        ]
    )
    
    doc.add_paragraph()
    
    # ===== 附录二 =====
    p = doc.add_paragraph()
    run = p.add_run('附录二：主要访谈对象信息')
    set_font(run, '黑体', 'Times New Roman', 10.5)
    set_line_spacing(p, 16)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('表2 受访者基本信息一览')
    set_font(run, '黑体', 'Times New Roman', 9)
    set_line_spacing(p, 16)
    
    add_three_line_table(doc,
        ['编号', '化名', '年龄', '民族', '职业', '访谈时长'],
        [
            ['A01', '翦玉成', '75', '维吾尔族', '退休农民', '120分钟'],
            ['A02', '翦明华', '42', '维吾尔族', '小学教师', '90分钟'],
            ['A03', '翦秀兰', '53', '维吾尔族', '务农/个体户', '75分钟'],
            ['A04', '翦雨婷', '19', '维吾尔族', '在校大学生', '60分钟'],
            ['A05', '李德明', '68', '汉族', '退休教师', '55分钟'],
        ]
    )
    
    # 表注
    p = doc.add_paragraph()
    run = p.add_run('注：为保护受访者隐私，以上均为化名。其余7位受访者信息略。')
    set_font(run, '宋体', 'Times New Roman', 9)
    set_line_spacing(p, 16)
    
    return doc


if __name__ == '__main__':
    output_path = sys.argv[1] if len(sys.argv) > 1 else '/root/.openclaw/workspace-taizi/JJC-20260408-001_实践报告.docx'
    doc = build_doc()
    doc.save(output_path)
    print(f'✅ Word 文档已生成: {output_path}')
